from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Proposal_Ghana_Chamber_of_Mines_BrandWatch_Africa.docx"

BLUE = RGBColor(31, 77, 120)
MID_BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(22, 34, 49)
MUTED = RGBColor(90, 100, 112)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "E8EEF5"
BORDER = "D9E2EC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=130, bottom=90, end=130):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def set_run(run, size=11, color=None, bold=None, italic=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_para(doc, text="", style=None, bold=False, color=None, size=11, after=6, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run(run, size=size, color=color, bold=bold)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.10
        run = p.add_run(item)
        set_run(run)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.10
        run = p.add_run(item)
        set_run(run)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    if level == 1:
        set_run(run, size=16, color=MID_BLUE, bold=True)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        set_run(run, size=13, color=MID_BLUE, bold=True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        set_run(run, size=12, color=BLUE, bold=True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_callout(doc, label, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    set_table_borders(table, "C8D4E3")
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    set_cell_margins(cell, top=130, bottom=130, start=170, end=170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label)
    set_run(r, size=10.5, color=BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run(r2, size=10.5, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_scope_table(doc):
    rows = [
        ("Option", "Scope", "Best For"),
        ("Chamber Sector Pulse", "Monitor industry-wide mining narratives across selected public sources", "Establishing the Chamber's baseline intelligence capability"),
        ("Corridor Pilot", "Focus on one mining corridor or region with selected member companies and community topics", "Testing community sentiment and local issue detection"),
        ("Member Consortium Pilot", "Include a small group of participating member companies with anonymized benchmark reporting", "Proving value for member expansion and paid rollout"),
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    set_table_width(table, [1.55, 2.75, 2.20])
    set_table_borders(table)
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_FILL)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            set_run(run, size=9.5 if r_idx else 10, color=DARK, bold=(r_idx == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_title_page(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.text = "BrandWatch Africa | Ghana Chamber of Mines Proposal"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "June 2026"
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED

    add_para(doc, "BrandWatch Africa / SIGNAL", bold=True, color=MUTED, size=12, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    title = add_para(
        doc,
        "Proposal to the Ghana Chamber of Mines",
        bold=True,
        color=DARK,
        size=24,
        after=4,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    title.paragraph_format.space_before = Pt(22)
    add_para(
        doc,
        "AI-Powered Reputation, Community Sentiment, and ESG Intelligence for Ghana's Mining Sector",
        color=MUTED,
        size=13.5,
        after=24,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    table = doc.add_table(rows=4, cols=2)
    set_table_width(table, [1.55, 4.95])
    set_table_borders(table, "E1E7EF")
    data = [
        ("Prepared for", "The Ghana Chamber of Mines"),
        ("Prepared by", "BrandWatch Africa / SIGNAL Reputation Intelligence Platform"),
        ("Date", "June 2026"),
        ("Proposed engagement", "Chamber-led pilot and sector intelligence partnership"),
    ]
    for row, (label, value) in zip(table.rows, data):
        for idx, text in enumerate([label, value]):
            cell = row.cells[idx]
            set_cell_margins(cell, top=110, bottom=110)
            if idx == 0:
                set_cell_shading(cell, LIGHT_FILL)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run(r, size=10.5, color=DARK, bold=(idx == 0))

    doc.add_paragraph().paragraph_format.space_after = Pt(18)
    add_callout(
        doc,
        "Proposed partnership",
        "An 8 to 12 week Chamber-led pilot to detect mining-sector reputation risks earlier, improve ESG evidence, and create a scalable sentiment intelligence capability for member companies.",
    )
    doc.add_page_break()


def main():
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Calibri"

    add_title_page(doc)

    add_heading(doc, "Executive Summary", 1)
    add_para(doc, "Ghana's mining industry operates in a high-trust, high-scrutiny environment. Member companies are expected to maintain strong community relationships, demonstrate responsible environmental and social performance, respond quickly to misinformation, and provide credible evidence of social license to operate. Yet many of the most important reputation signals are fragmented across radio discussions, local-language conversations, community reports, online news, social media, stakeholder meetings, and field officer notes.")
    add_para(doc, "BrandWatch Africa is proposed as a Chamber-led reputation and community sentiment intelligence platform for the mining sector. The platform combines multi-channel monitoring, Ghanaian and African language sentiment analysis, stakeholder-specific risk scoring, narrative tracking, and ESG-ready reporting.")
    add_callout(doc, "Core promise", "Move from simply knowing that mining is being discussed to knowing which communities are losing trust, which issues are driving sentiment, where narratives are spreading, and what action is needed before concerns escalate.")
    add_para(doc, "We recommend an initial 8 to 12 week pilot focused on one or more mining corridors, selected member companies, and sector-wide narratives. The pilot would produce a live intelligence dashboard, weekly insight briefings, issue and risk maps, a final executive report, and a clear roadmap for expansion.")

    add_heading(doc, "Why This Matters Now", 1)
    add_para(doc, "Mining companies in Ghana face social, regulatory, investor, and communications pressures. Community expectations around participation, compensation, jobs, environmental protection, local content, and sustainable development continue to rise. ESG reporting expectations require stronger evidence, not only anecdotes.")
    add_para(doc, "Existing global social listening tools are not built for this context. They often emphasize online mentions, English-language social media, and generic sentiment labels. They can miss the channels and language patterns that matter most in Ghanaian mining communities, including local radio, community reports, traditional stakeholder concerns, code-switching, Ghanaian English, Twi, Akan, Ga, Ewe, Dagbani, Hausa, and French-language regional content.")
    add_para(doc, "For the Chamber, this creates an opportunity to lead. A sector-level reputation intelligence system can help the Chamber understand the health of the industry's public standing, detect emerging issues early, support member companies with better insight, and strengthen engagement with government, regulators, host communities, investors, media, and civil society.")

    add_heading(doc, "Proposed Solution", 1)
    add_para(doc, "BrandWatch Africa is an AI-powered brand identity and reputation monitoring platform built for African markets and high-impact sectors. For the Ghana Chamber of Mines, the platform would be configured as a mining-sector intelligence system with four connected capabilities.")
    add_heading(doc, "1. Multi-Channel Monitoring", 2)
    add_bullets(doc, [
        "Online news, blogs, and public web sources",
        "Social media posts and comments where compliant access is available",
        "Radio transcripts and audio-derived text where available",
        "Public regulator and government statements",
        "Uploaded field reports, survey exports, call center logs, or community relations notes",
        "Stakeholder and community feedback datasets provided by member companies",
    ])
    add_heading(doc, "2. Local Context Sentiment and Narrative Intelligence", 2)
    add_para(doc, "The sentiment engine classifies mentions as positive, negative, neutral, mixed, or uncertain while accounting for Ghanaian and wider African communication patterns. It supports local-language sentiment detection, code-switching, sarcasm, indirect criticism, praise-with-complaint, community grievance language, and culturally specific expressions.")
    add_heading(doc, "3. Community and Sector Risk Scoring", 2)
    add_para(doc, "BrandWatch Africa converts raw mentions into prioritized risk signals. Each issue can receive a risk score based on sentiment intensity, mention volume change, source credibility, spread velocity, topic severity, location, stakeholder group, and whether similar issues have previously escalated.")
    add_heading(doc, "4. ESG and Executive Reporting", 2)
    add_bullets(doc, [
        "Industry reputation trend reports",
        "Community trust and social license indicators",
        "Sentiment by region, site, channel, topic, and stakeholder group",
        "Before-and-after analysis for engagement initiatives",
        "Issue heatmaps and narrative timelines",
        "Executive-ready summaries for board, ESG, and communications use",
    ])

    add_heading(doc, "Value to the Ghana Chamber of Mines", 1)
    add_heading(doc, "Sector Reputation Management", 2)
    add_para(doc, "The Chamber would gain a clearer view of how mining is being discussed across media, digital, and community channels. This supports proactive communication, more precise public education, and stronger response to narratives that may affect the whole sector.")
    add_heading(doc, "Early Warning on Emerging Issues", 2)
    add_para(doc, "The system can surface rising concerns before they become national headlines, protests, regulatory scrutiny, or reputational crises. It can help the Chamber identify when a local issue is becoming an industry issue.")
    add_heading(doc, "Better Member Support", 2)
    add_para(doc, "The Chamber can provide members with evidence-based insight into sentiment trends, recurring concerns, and stakeholder priorities. Over time, anonymized benchmarking can help members compare community trust, issue volume, response times, and improvement trends.")
    add_heading(doc, "ESG and Social License Evidence", 2)
    add_para(doc, "The platform gives the Chamber and its members measurable evidence for ESG reporting, investor engagement, sustainability narratives, and responsible mining commitments. It can help translate community perception into a trackable social license indicator.")

    add_heading(doc, "Value to Member Companies", 1)
    add_bullets(doc, [
        "Detect community grievances before they escalate",
        "Understand which topics drive trust or distrust around specific sites",
        "Track sentiment around jobs, compensation, local content, environment, safety, CSR, and land issues",
        "Reduce manual time spent compiling media and community reports",
        "Support ESG, board, and investor reporting with more defensible data",
        "Measure whether engagement activities improve sentiment over time",
        "Benchmark company reputation against sector trends and peer performance",
    ])

    add_heading(doc, "Proposed Pilot", 1)
    add_para(doc, "We recommend an 8 to 12 week pilot designed to prove value quickly while respecting data access, privacy, member confidentiality, and the Chamber's coordination role.")
    add_heading(doc, "Pilot Objectives", 2)
    add_numbers(doc, [
        "Identify dominant narratives affecting mining-sector reputation in selected Ghanaian markets.",
        "Understand which topics, channels, communities, and stakeholder groups drive positive and negative sentiment.",
        "Detect emerging issues with early signs of escalation.",
        "Test how sentiment intelligence can support ESG reporting and stakeholder engagement.",
        "Define the data sources, workflows, and governance model needed for sector rollout.",
    ])
    add_heading(doc, "Suggested Scope", 2)
    add_scope_table(doc)
    add_heading(doc, "Pilot Deliverables", 2)
    add_bullets(doc, [
        "Configured BrandWatch Africa dashboard for Chamber users",
        "Mining-sector topic taxonomy covering environment, employment, compensation, safety, local content, community investment, illegal mining association risk, regulation, media narratives, and stakeholder trust",
        "Weekly intelligence briefs with top narratives, sentiment movement, emerging risks, and recommended actions",
        "Risk queue prioritized by urgency, location, stakeholder group, and topic",
        "Human review workflow for low-confidence or culturally ambiguous sentiment",
        "Final executive report with findings, sample dashboards, pilot metrics, and scale-up recommendations",
    ])

    add_heading(doc, "Implementation Roadmap", 1)
    roadmap = [
        ("Phase 1: Sentiment MVP", "Ingestion from CSV uploads, news feeds, and selected social sources; basic language detection; sentiment labels; confidence scoring; topic tags; a basic dashboard; and manual review."),
        ("Phase 2: Context-Aware Sentiment", "Local-language examples, code-switching support, emotion and intent detection, explainable sentiment output, stakeholder classification, location extraction, and sentiment views by community, region, source, and stakeholder group."),
        ("Phase 3: Risk and Alerts", "Community risk scoring, alert thresholds, narrative velocity tracking, suggested response owners, escalation workflows, and alerts through email, Slack, SMS, or WhatsApp Business where appropriate."),
        ("Phase 4: ESG and Executive Reporting", "ESG-aligned report templates, social license trend score, community trust index, board-ready exports, competitor benchmarking, intervention impact analysis, and data export API."),
        ("Phase 5: Enterprise Intelligence Platform", "Multi-country support, custom model tuning, advanced radio and audio transcription, narrative clustering, bot and synthetic content detection, predictive crisis forecasting, enterprise permissions, audit logs, and managed analyst workflows."),
    ]
    for label, body in roadmap:
        add_heading(doc, label, 2)
        add_para(doc, body)

    add_heading(doc, "Governance and Data Protection", 1)
    add_bullets(doc, [
        "Use public, licensed, customer-provided, or ethically collected data sources",
        "Minimize personal data collection wherever possible",
        "Redact personally identifiable information in field reports and call center exports where appropriate",
        "Define role-based access for Chamber users, participating members, and analysts",
        "Separate company-specific views from anonymized sector-wide benchmark views",
        "Route low-confidence or culturally sensitive classifications for human review",
        "Agree on a protocol for urgent issue escalation and member notification",
    ])

    add_heading(doc, "Success Metrics", 1)
    add_heading(doc, "Product Performance Metrics", 2)
    add_bullets(doc, [
        "Sentiment classification agreement with human reviewers",
        "Percentage of mentions classified with high confidence",
        "Time from mention ingestion to usable insight",
        "Number of high-risk issues detected",
        "Alert precision and usefulness",
        "Dashboard usage by Chamber and pilot stakeholders",
    ])
    add_heading(doc, "Business and Stakeholder Outcome Metrics", 2)
    add_bullets(doc, [
        "Earlier detection of community or media issues",
        "Reduction in manual reporting time",
        "Quality of weekly insight briefings",
        "Usefulness of ESG-ready evidence",
        "Ability to identify dominant narratives and issue drivers",
        "Member appetite for expansion after pilot completion",
    ])

    add_heading(doc, "Commercial Model", 1)
    add_para(doc, "We recommend a pilot-first commercial approach. The pilot can be structured as a fixed-scope engagement covering setup, taxonomy design, data ingestion, dashboard configuration, weekly analysis, and final reporting. Final pricing should be determined after confirming source volume, number of users, member participation, and whether radio transcription is included.")
    add_bullets(doc, [
        "Chamber subscription for sector-wide reputation intelligence",
        "Member company subscriptions for company-specific dashboards",
        "Consortium model with participating members sharing base platform costs",
        "Managed analyst service for monthly or quarterly insight reviews",
        "Add-on modules for radio monitoring, ESG report automation, crisis dashboards, historical backfill, custom board reports, and API access",
    ])

    add_heading(doc, "Partnership Ask", 1)
    add_para(doc, "We propose that the Ghana Chamber of Mines partner with BrandWatch Africa to design and run a mining-sector reputation intelligence pilot.")
    add_numbers(doc, [
        "Nominate a Chamber pilot sponsor and working group.",
        "Select the pilot model: Sector Pulse, Corridor Pilot, or Member Consortium Pilot.",
        "Agree on initial keywords, member companies, regions, topics, and reporting needs.",
        "Confirm available data sources and governance requirements.",
        "Run a 2 week setup sprint followed by 8 to 12 weeks of monitoring and reporting.",
        "Review pilot outcomes and decide on a sector rollout model.",
    ])

    add_heading(doc, "Conclusion", 1)
    add_para(doc, "The Ghana Chamber of Mines is well positioned to lead a new standard for reputation intelligence and evidence-based stakeholder engagement in the mining sector. BrandWatch Africa can support that leadership by giving the Chamber and its members a practical system for listening earlier, understanding better, prioritizing risk, and reporting with confidence.")
    add_para(doc, "The result is not just more monitoring. It is a stronger operating capability for responsible mining: one that connects community sentiment, media narratives, ESG evidence, and executive action in a single, Ghana-ready intelligence platform.")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
